"""
Market Intelligence → DOCX Exporter (Robust Version)
=====================================================
- LaTeX inline ($...$) dan display ($$...$$) dirender ke Office Math (OMML)
- Mermaid charts dirender ke gambar via mermaid.ink
- Fallback aman: jika OMML gagal → teks unicode bersih (tidak crash)
- Tabel Markdown → Word Table
- Cover page profesional
"""

import os
import re
import base64
import requests
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml import parse_xml as docx_parse_xml
from docx.oxml.ns import nsdecls, qn

try:
    import lxml.etree as ET
    HAS_LXML = True
except ImportError:
    HAS_LXML = False

try:
    import latex2mathml.converter
    HAS_LATEX2MATHML = True
except ImportError:
    HAS_LATEX2MATHML = False

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
XSL_PATH    = os.path.join(SCRIPT_DIR, "mml2omml.xsl")
OUTPUT_PATH = r"D:\txt\Sistem_Intelijen_Pasar_2_0_Lengkap.docx"

MML2OMML_MIRRORS = [
    "https://raw.githubusercontent.com/MartinPaulEve/meTypeset/master/docx/utils/maths/mml2omml.xsl",
    "https://raw.githubusercontent.com/nicktindall/cyclo/main/mml2omml.xsl",
    "https://raw.githubusercontent.com/w3c/mml-xslt/master/mml2omml.xsl",
]

M_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"m": M_NS, "w": W_NS}

# ─────────────────────────────────────────────────────────────────────────────
# DOWNLOAD XSL IF MISSING & POST-PROCESS
# ─────────────────────────────────────────────────────────────────────────────
def ensure_xsl():
    downloaded = False
    if not os.path.exists(XSL_PATH) or os.path.getsize(XSL_PATH) < 1000:
        for url in MML2OMML_MIRRORS:
            try:
                print(f"[XSL] Downloading from {url} ...")
                r = requests.get(url, timeout=15)
                if r.status_code == 200:
                    with open(XSL_PATH, "wb") as f:
                        f.write(r.content)
                    print(f"[XSL] Saved to {XSL_PATH}")
                    downloaded = True
                    break
            except Exception as e:
                print(f"[XSL] Failed to download: {e}")
    else:
        downloaded = True

    if downloaded or os.path.exists(XSL_PATH):
        try:
            with open(XSL_PATH, "r", encoding="utf-8") as f:
                content = f.read()
            
            # Uncomment the root template match="/" if it is commented out
            pattern = r'<!--\s*(<xsl:template\s+match="/".*?</xsl:template>)\s*-->'
            if re.search(pattern, content, re.DOTALL):
                print("[XSL] Uncommenting the root template in mml2omml.xsl...")
                uncommented = re.sub(pattern, r'\1', content, flags=re.DOTALL)
                with open(XSL_PATH, "w", encoding="utf-8") as f:
                    f.write(uncommented)
                print("[XSL] mml2omml.xsl prepared successfully.")
            return True
        except Exception as e:
            print(f"[XSL] Error post-processing mml2omml.xsl: {e}")
    return False

HAS_XSL = ensure_xsl()

# ─────────────────────────────────────────────────────────────────────────────
# LATEX CLEANING
# ─────────────────────────────────────────────────────────────────────────────
def strip_dollar(s: str) -> str:
    s = s.strip()
    if s.startswith("$$") and s.endswith("$$"):
        return s[2:-2].strip()
    if s.startswith("$") and s.endswith("$"):
        return s[1:-1].strip()
    return s

def clean_latex(latex: str) -> str:
    """Minimal cleaning – preserve structure for latex2mathml."""
    t = strip_dollar(latex)
    # Replace escaped underscores inside \text{} with a space
    def fix_text_block(m):
        inner = m.group(1).replace(r'\_', ' ').replace('_', ' ')
        return r'\text{' + inner + '}'
    t = re.sub(r'\\text\{([^}]*)\}', fix_text_block, t)
    return t

# ─────────────────────────────────────────────────────────────────────────────
# UNICODE FALLBACK RENDERER
# ─────────────────────────────────────────────────────────────────────────────
UNICODE_MAP = [
    (r'\\sigma', 'σ'), (r'\\alpha', 'α'), (r'\\beta', 'β'), (r'\\gamma', 'γ'),
    (r'\\delta', 'δ'), (r'\\epsilon', 'ε'), (r'\\mu', 'μ'), (r'\\rho', 'ρ'),
    (r'\\pi', 'π'), (r'\\theta', 'θ'), (r'\\lambda', 'λ'), (r'\\tau', 'τ'),
    (r'\\Sigma', 'Σ'), (r'\\sqrt', '√'), (r'\\times', '×'), (r'\\cdot', '·'),
    (r'\\le', '≤'), (r'\\ge', '≥'), (r'\\neq', '≠'), (r'\\approx', '≈'),
    (r'\\infty', '∞'), (r'\\sum', 'Σ'), (r'\\prod', 'Π'), (r'\\ln', 'ln'),
    (r'\\log', 'log'), (r'\\max', 'max'), (r'\\min', 'min'), (r'\\quad', '  '),
    (r'\\left', ''), (r'\\right', ''), (r'\\hat', ''), (r'\\bar', ''),
    (r'\\begin\{cases\}', '['), (r'\\end\{cases\}', ']'),
]

def latex_to_unicode(latex: str) -> str:
    t = clean_latex(latex)
    # \frac{a}{b} → (a)/(b)
    def replace_frac(m): return f"({m.group(1)})/({m.group(2)})"
    t = re.sub(r'\\frac\s*\{([^{}]*)\}\s*\{([^{}]*)\}', replace_frac, t)
    # \text{...} → its content
    t = re.sub(r'\\text\{([^}]*)\}', r'\1', t)
    # ^ subscript/superscript
    t = re.sub(r'\^\\?\{?([^}]+)\}?', r'^\1', t)
    t = re.sub(r'_\\?\{?([^}]+)\}?', r'_\1', t)
    for pattern, repl in UNICODE_MAP:
        t = re.sub(pattern, repl, t)
    # strip remaining backslashes
    t = re.sub(r'\\[a-zA-Z]+', '', t)
    t = t.replace('\\', '').replace('{', '').replace('}', '')
    # remove double spaces
    t = re.sub(r'  +', ' ', t).strip()
    return t

# ─────────────────────────────────────────────────────────────────────────────
# OMML BUILDER (LaTeX → MathML → OMML via XSLT)
# ─────────────────────────────────────────────────────────────────────────────
_xsl_transform = None

def get_xslt_transform():
    global _xsl_transform
    if _xsl_transform is None and HAS_LXML and HAS_XSL:
        try:
            xslt_doc = ET.parse(XSL_PATH)
            _xsl_transform = ET.XSLT(xslt_doc)
        except Exception as e:
            print(f"[XSL] Failed to parse XSLT: {e}")
    return _xsl_transform

def make_omml_element(latex: str):
    """
    Returns an lxml/docx XML element: m:oMath or None on failure.
    Pipeline: LaTeX → MathML (latex2mathml) → OMML (XSLT) → lxml element
    """
    if not (HAS_LXML and HAS_LATEX2MATHML):
        return None

    cleaned = clean_latex(latex)
    transform = get_xslt_transform()
    if transform is None:
        return None

    try:
        mathml_str = latex2mathml.converter.convert(cleaned)
        mathml_xml  = ET.fromstring(mathml_str.encode("utf-8"))
        omml_tree   = transform(mathml_xml)
        omml_bytes  = ET.tostring(omml_tree, encoding="utf-8")
        omml_str    = omml_bytes.decode("utf-8")
        # Strip XML declaration
        omml_str = re.sub(r'<\?xml[^>]*\?>', '', omml_str).strip()
        if not omml_str:
            return None

        element = docx_parse_xml(omml_str)

        # If wrapped in oMathPara, extract inner oMath
        tag_local = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag_local == 'oMathPara':
            inner = element.find(f'{{{M_NS}}}oMath')
            if inner is not None:
                return inner
            # Also try without namespace
            inner = element.find('.//oMath')
            if inner is not None:
                return inner

        return element

    except Exception as e:
        print(f"[MATH] Failed to convert LaTeX '{latex[:60]}...': {e}")
        return None


def make_fallback_omml(text: str):
    """Build a minimal m:oMath element with plain text as last resort."""
    safe = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
    xml = (
        f'<m:oMath xmlns:m="{M_NS}" xmlns:w="{W_NS}">'
        f'<m:r><m:t>{safe}</m:t></m:r>'
        f'</m:oMath>'
    )
    try:
        return docx_parse_xml(xml)
    except Exception:
        return None


def append_math(paragraph, latex: str, is_display: bool = False):
    """
    Append math to paragraph. 
    First try OMML pipeline, fall back to unicode text.
    """
    omml = make_omml_element(latex)
    if omml is not None:
        paragraph._p.append(omml)
        return True
    # Fallback unicode
    unicode_text = latex_to_unicode(latex)
    fb_omml = make_fallback_omml(unicode_text)
    if fb_omml is not None:
        paragraph._p.append(fb_omml)
        return True
    # Last resort: plain run
    run = paragraph.add_run(unicode_text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return False

# ─────────────────────────────────────────────────────────────────────────────
# MERMAID → IMAGE
# ─────────────────────────────────────────────────────────────────────────────
def render_mermaid(code: str) -> str | None:
    try:
        # Strip init directives that mermaid.ink doesn't support
        clean = re.sub(r'^%%\{.*?\}%%\s*', '', code.strip(), flags=re.DOTALL)
        b64   = base64.urlsafe_b64encode(clean.encode("utf-8")).decode("utf-8")
        url   = f"https://mermaid.ink/img/{b64}?type=png"
        r     = requests.get(url, timeout=20)
        if r.status_code == 200 and len(r.content) > 500:
            tmp = os.path.join(SCRIPT_DIR, "_mermaid_tmp.png")
            with open(tmp, "wb") as f:
                f.write(r.content)
            return tmp
    except Exception as e:
        print(f"[MERMAID] Render failed: {e}")
    return None

# ─────────────────────────────────────────────────────────────────────────────
# INLINE PARSER: handles $...$ and **...**
# ─────────────────────────────────────────────────────────────────────────────
# Split on: $...$ (non-greedy, single-line) OR **...**
INLINE_SPLIT = re.compile(r'(\$[^$\n]+?\$|\*\*[^*]+?\*\*|`[^`]+?`)')

def parse_inline(paragraph, text: str):
    tokens = INLINE_SPLIT.split(text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('$') and tok.endswith('$') and len(tok) > 2:
            append_math(paragraph, tok, is_display=False)
        elif tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(tok)

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: cell margins
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# ─────────────────────────────────────────────────────────────────────────────
# CALLOUT BOX
# ─────────────────────────────────────────────────────────────────────────────
def add_callout(doc, text: str):
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shd  = docx_parse_xml(r'<w:shd {} w:fill="EFF6FF"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)
    borders = docx_parse_xml(r'''
        <w:tcBorders {}>
            <w:top    w:val="none"/>
            <w:left   w:val="single" w:sz="24" w:space="0" w:color="3B82F6"/>
            <w:bottom w:val="none"/>
            <w:right  w:val="none"/>
        </w:tcBorders>'''.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(borders)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    parse_inline(p, text.strip())
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# HEADING STYLES
# ─────────────────────────────────────────────────────────────────────────────
def add_heading(doc, text: str, level: int):
    colours = {1: RGBColor(0x1A, 0x36, 0x5D),
               2: RGBColor(0x2C, 0x52, 0x82),
               3: RGBColor(0x4A, 0x55, 0x68)}
    sizes   = {1: 16, 2: 13, 3: 11.5}
    before  = {1: 22, 2: 16, 3: 12}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before.get(level, 10))
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.strip())
    run.bold = True
    run.font.size  = Pt(sizes.get(level, 11))
    run.font.color.rgb = colours.get(level, RGBColor(0, 0, 0))

# ─────────────────────────────────────────────────────────────────────────────
# TABLE RENDERER
# ─────────────────────────────────────────────────────────────────────────────
def flush_table(doc, rows_data: list):
    """Render accumulated Markdown table rows into a Word table."""
    # Filter separator rows (---|--- patterns)
    rows_data = [r for r in rows_data if not re.match(r'^\|\s*[-:]+[\s|:-]*\|?\s*$', r)]
    if not rows_data:
        return

    parsed = []
    for row in rows_data:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    max_cols = max(len(r) for r in parsed)
    word_tbl = doc.add_table(rows=len(parsed), cols=max_cols)
    word_tbl.style = 'Table Grid'

    for r_i, row in enumerate(parsed):
        for c_i in range(max_cols):
            cell_text = row[c_i] if c_i < len(row) else ''
            cell = word_tbl.cell(r_i, c_i)
            p    = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            parse_inline(p, cell_text)
            set_cell_margins(cell)
            if r_i == 0:
                shd = docx_parse_xml(r'<w:shd {} w:fill="DBEAFE"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# CORE FILE PROCESSOR
# ─────────────────────────────────────────────────────────────────────────────
def process_md_file(doc, filepath: str):
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    # ── Split on display math blocks $$ ... $$ (multiline) ──
    # We'll process the file line-by-line with a state machine instead, 
    # which is more robust
    lines = content.split("\n")

    in_code      = False
    is_mermaid   = False
    code_buf     = []
    in_display   = False
    display_buf  = []
    table_buf    = []

    def flush_table_buf():
        if table_buf:
            flush_table(doc, list(table_buf))
            table_buf.clear()

    for line in lines:
        # ── Display math: $$ ... $$ (possibly multi-line) ──
        if not in_code and not in_display:
            stripped = line.strip()
            if stripped == '$$' or stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
                flush_table_buf()
                if stripped == '$$':
                    in_display = True
                    display_buf = []
                else:
                    # Entire display block on one line
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after  = Pt(10)
                    append_math(p, stripped, is_display=True)
                continue
            elif re.match(r'^\$\$\s*.+', stripped) and '$$' not in stripped[2:]:
                # $$ starts on this line but doesn't end here
                flush_table_buf()
                in_display = True
                display_buf = [stripped[2:]]
                continue

        if in_display:
            if line.strip() == '$$' or line.strip().endswith('$$'):
                # End of display block
                remainder = line.strip()
                if remainder != '$$':
                    display_buf.append(remainder.rstrip('$').rstrip())
                full_latex = ' '.join(display_buf).strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(10)
                append_math(p, full_latex, is_display=True)
                in_display  = False
                display_buf = []
            else:
                display_buf.append(line)
            continue

        # ── Code blocks ──
        if line.strip().startswith("```"):
            flush_table_buf()
            if not in_code:
                in_code    = True
                is_mermaid = 'mermaid' in line.lower()
                code_buf   = []
            else:
                in_code = False
                code_text = "\n".join(code_buf)
                if is_mermaid:
                    img = render_mermaid(code_text)
                    if img and os.path.exists(img):
                        pi = doc.add_paragraph()
                        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pi.add_run().add_picture(img, width=Inches(6.0))
                        os.remove(img)
                    else:
                        p = doc.add_paragraph()
                        p.add_run("[Mind Map / Diagram – visualisasikan di browser GitHub]").italic = True
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run(code_text.rstrip())
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Tables ──
        if line.strip().startswith("|"):
            table_buf.append(line.strip())
            continue
        else:
            flush_table_buf()

        # ── Empty line ──
        stripped = line.strip()
        if stripped == "":
            continue

        # ── Headings ──
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:], 3)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:], 2)
        elif stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)

        # ── Blockquote / callout ──
        elif stripped.startswith(">"):
            inner = re.sub(r'^>+\s*', '', stripped)
            # Skip GitHub alert markers like [!NOTE], [!TIP] etc.
            if re.match(r'^\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', inner):
                pass
            elif inner:
                add_callout(doc, inner)

        # ── Horizontal rule ──
        elif re.match(r'^[-*_]{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            pPr = OxmlElement('w:pPr')
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CBD5E1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            p._p.insert(0, pPr)

        # ── List items ──
        elif re.match(r'^[-*+]\s', stripped) or re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(3)
            # Remove list marker
            text_content = re.sub(r'^[-*+]\s+|^\d+\.\s+', '', stripped)
            parse_inline(p, text_content)

        # ── Normal paragraph ──
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parse_inline(p, stripped)

    flush_table_buf()

# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
SOURCE_FILES = [
    {"path": r"D:\Market-Intelligence\README.md",
     "desc": "PANDUAN UTAMA & README SISTEM"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\01_General_Overview.md",
     "desc": "PART 1 – SYSTEM ONTOLOGY, CAUSAL HIERARCHY & HIGH-LEVEL ARCHITECTURE"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\02_Data_Pipeline_and_Security.md",
     "desc": "PART 2 – DATA ENGINEERING, DATABASE INFRASTRUCTURE & CAUSAL INTEGRITY"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\03_Macroeconomic_and_Quant_Finance.md",
     "desc": "PART 3 – MACROECONOMIC REGIMES & QUANTITATIVE FEATURE ENGINEERING"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\04_Machine_Learning_Architecture.md",
     "desc": "PART 4 – ENSEMBLE MODELING & DUAL-HEAD META-LEARNING ARCHITECTURE"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\05_Workflow_and_Inference_Engine.md",
     "desc": "PART 5 – PRODUCTION INFERENCE ENGINE & TRAJECTORY SMOOTHING"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\06_Correlation_Enforcer_and_Beta_Correction.md",
     "desc": "PART 6 – INTER-ASSET CORRELATION & BETA-ADJUSTED FORECAST CORRECTOR"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\07_XAI_Explainer_and_ZScore_Attribution.md",
     "desc": "PART 7 – EXPLAINABLE AI, FREQUENCY-AWARE Z-SCORE & LLM ATTRIBUTION"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\08_Trading_Signal_Generator.md",
     "desc": "PART 8 – MULTI-FACTOR TRADING SIGNAL GENERATION & RISK CALIBRATION"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\09_Counterfactual_Logging_and_Model_Validation.md",
     "desc": "PART 9 – COUNTERFACTUAL BACKTESTING & MODEL ACCURACY SCORES"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\10_Alternative_Data_Ingestion.md",
     "desc": "PART 10 – ALTERNATIVE DATA INGESTION PIPELINES & API BINDINGS"},
    {"path": r"D:\Market-Intelligence\docs\system_architecture\11_Critiques_Defenses_and_OOD_Anomaly_Gate.md",
     "desc": "PART 11 – CRITIQUES, ARCHITECTURAL DEFENSES & OOD ANOMALY GATE"},
]


def build_docx():
    doc = Document()

    # ── Page setup ──
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11.0)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    # ── Base style ──
    normal = doc.styles['Normal']
    normal.font.name  = 'Arial'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after  = Pt(6)

    # ── Cover Page ──
    doc.add_paragraph().paragraph_format.space_before = Pt(100)

    p = doc.add_paragraph()
    r = p.add_run("MARKET INTELLIGENCE SYSTEM")
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Sistem Intelijen Pasar 2.0")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Buku Putih Riset & Spesifikasi Arsitektur Teknis Komprehensif")
    r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    p3.paragraph_format.space_after = Pt(120)

    p4 = doc.add_paragraph()
    r4 = p4.add_run(
        "Klasifikasi : Dokumen Teknis Institusional (Production-Grade)\n"
        "Pipeline    : LaTeX → MathML → OMML (Office Math)\n"
        "Sumber      : README + 11 Dokumen Arsitektur Sistem\n"
        "Tanggal     : Mei 2026"
    )
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    doc.add_page_break()

    # ── Process each source ──
    for item in SOURCE_FILES:
        path = item["path"]
        desc = item["desc"]

        if not os.path.exists(path):
            print(f"[SKIP] File not found: {path}")
            continue

        print(f"[OK]   Processing: {os.path.basename(path)}")

        # Section divider
        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_before = Pt(20)
        p_div.paragraph_format.space_after  = Pt(12)
        run = p_div.add_run(f"{'─' * 6}  {desc}  {'─' * 6}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        process_md_file(doc, path)
        doc.add_page_break()

    # ── Save ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"\n[DONE] Saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_docx()
