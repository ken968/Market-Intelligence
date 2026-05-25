import docx

doc_path = r"D:\txt\Sistem_Intelijen_Pasar_2_0_Lengkap.docx"
try:
    doc = docx.Document(doc_path)
    print("Document loaded successfully.")
    
    # Check for placeholder text
    placeholders = []
    for p_idx, p in enumerate(doc.paragraphs):
        if "[Mind Map" in p.text or "visualisasikan di browser" in p.text:
            placeholders.append((p_idx, p.text))
            
    print(f"Found {len(placeholders)} placeholders:")
    for idx, text in placeholders:
        print(f"  Paragraph {idx}: {text}")
        
    # Check inline math fallbacks
    fallbacks = []
    for p_idx, p in enumerate(doc.paragraphs):
        if "Fallback" in p.text or "fallback" in p.text:
            pass # just checking
            
    # Check how many images are in the document
    # Images are in doc.inline_shapes
    print(f"Number of inline shapes: {len(doc.inline_shapes)}")
    
except Exception as e:
    print("Error:", e)
