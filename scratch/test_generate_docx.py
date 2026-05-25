import docx
from docx import Document
from docx.oxml import parse_xml as docx_parse_xml
import lxml.etree as ET
import latex2mathml.converter
import re

XSL_PATH = r"d:\Market-Intelligence\scripts\mml2omml.xsl"
xslt_doc = ET.parse(XSL_PATH)
transform = ET.XSLT(xslt_doc)

doc = Document()
p = doc.add_paragraph()
p.add_run("Here is some math: ")

def add_math_to_p(paragraph, latex):
    mathml_str = latex2mathml.converter.convert(latex)
    mathml_xml = ET.fromstring(mathml_str.encode("utf-8"))
    omml_tree = transform(mathml_xml)
    omml_bytes = ET.tostring(omml_tree, encoding="utf-8")
    omml_str = omml_bytes.decode("utf-8")
    omml_str = re.sub(r'<\?xml[^>]*\?>', '', omml_str).strip()
    
    # DO NOT replace <oMath> with <m:oMath>!
    element = docx_parse_xml(omml_str)
    paragraph._p.append(element)

add_math_to_p(p, r"\rho")
p.add_run(" and another: ")
add_math_to_p(p, r"y_t = \frac{\text{Price}_{t + H} - \text{Price}_t}{\text{Price}_t}")

doc.save(r"d:\Market-Intelligence\scratch\test_math_out.docx")
print("Saved successfully!")
